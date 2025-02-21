"""
Word文档解析器
"""
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import docx
from base.core.logging import LogLevel
from base.core.logging import AsyncLogger
from base.rag.document import (Document, DocumentMetadata, DocumentParser,
                             DocumentType)
from module.documents import TextDocument


class WordParser(DocumentParser):
    """Word文档解析器，支持.docx文件的解析"""
    
    def __init__(self, logger: Optional[AsyncLogger] = None):
        """初始化解析器
        
        Args:
            logger: 可选的日志记录器
        """
        super().__init__(logger)
    
    async def parse(self, file_path: Path, doc_type: DocumentType) -> Document:
        """解析单个Word文档
        
        Args:
            file_path: Word文件路径
            doc_type: 文档类型，必须是 DocumentType.WORD
            
        Returns:
            解析后的文档对象
            
        Raises:
            ValueError: 当文档类型不是WORD时抛出
        """
        await self._log(LogLevel.INFO, f"开始解析Word文档: {file_path}")
        
        if doc_type != DocumentType.WORD:
            await self._log(LogLevel.ERROR, f"文档类型不匹配: {doc_type}")
            raise ValueError(f"WordParser只支持WORD类型文档，收到: {doc_type}")
            
        if not file_path.exists():
            await self._log(LogLevel.ERROR, f"文件不存在: {file_path}")
            raise FileNotFoundError(f"文件不存在: {file_path}")
            
        try:
            # 打开Word文档
            await self._log(LogLevel.DEBUG, "打开Word文档")
            doc = docx.Document(file_path)
            
            # 提取文本内容
            content = []
            
            # 提取段落文本
            await self._log(LogLevel.DEBUG, f"开始处理段落，共 {len(doc.paragraphs)} 个段落")
            for i, paragraph in enumerate(doc.paragraphs, 1):
                text = paragraph.text.strip()
                if text:  # 只添加非空段落
                    # 如果当前段落是标题样式，添加额外的换行
                    if paragraph.style.name.startswith('Heading'):
                        if content:  # 不是第一个段落时添加额外换行
                            content.append('')
                        content.append(text)
                        content.append('')
                    else:
                        content.append(text)
                    await self._log(LogLevel.DEBUG, f"处理第 {i} 个段落")
            
            # 提取表格文本
            await self._log(LogLevel.DEBUG, f"开始处理表格，共 {len(doc.tables)} 个表格")
            for i, table in enumerate(doc.tables, 1):
                await self._log(LogLevel.DEBUG, f"处理第 {i} 个表格")
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:  # 只添加非空单元格
                            row_text.append(cell_text)
                    if row_text:  # 只添加非空行
                        table_content.append(" | ".join(row_text))
                if table_content:  # 如果表格有内容，添加到文档
                    if content:  # 不是第一个内容时添加空行
                        content.append('')
                    content.extend(table_content)
                    content.append('')
            
            # 合并所有文本，使用单个换行符分隔
            # 移除连续的空行，最多保留一个
            cleaned_content = []
            prev_empty = False
            for line in content:
                if line.strip():
                    cleaned_content.append(line)
                    prev_empty = False
                elif not prev_empty:
                    cleaned_content.append(line)
                    prev_empty = True
            
            full_content = '\n'.join(cleaned_content)
            
            # 创建元数据
            metadata = DocumentMetadata(
                doc_id=file_path.stem,
                doc_type=DocumentType.WORD,
                source=str(file_path),
                created_at=datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
                extra={
                    'file_size': file_path.stat().st_size,
                    'last_modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat(),
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables)
                }
            )
            
            await self._log(LogLevel.INFO, f"Word文档解析完成: {file_path}")
            return TextDocument(full_content, metadata)
            
        except Exception as e:
            await self._log(LogLevel.ERROR, f"解析Word文档时出错: {str(e)}")
            raise RuntimeError(f"解析Word文档时出错: {str(e)}")
    
    async def parse_batch(
        self,
        file_paths: List[Path],
        doc_types: List[DocumentType]
    ) -> List[Document]:
        """批量解析Word文档
        
        Args:
            file_paths: Word文件路径列表
            doc_types: 文档类型列表，必须全部是WORD类型
            
        Returns:
            解析后的文档对象列表
        """
        await self._log(LogLevel.INFO, f"开始批量解析 {len(file_paths)} 个Word文档")
        
        if not all(dt == DocumentType.WORD for dt in doc_types):
            await self._log(LogLevel.ERROR, "批量解析时发现不支持的文档类型")
            raise ValueError("WordParser只支持WORD类型文档")
            
        documents = []
        for file_path in file_paths:
            try:
                doc = await self.parse(file_path, DocumentType.WORD)
                documents.append(doc)
            except Exception as e:
                await self._log(LogLevel.ERROR, f"解析Word文档 {file_path} 时出错: {str(e)}")
                continue
                
        await self._log(LogLevel.INFO, f"批量解析完成，成功解析 {len(documents)} 个文件")
        return documents 